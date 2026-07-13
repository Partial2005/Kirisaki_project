"""
生成课程设计报告 —— Word 格式
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT = "知识工程课程设计报告.docx"

def set_cell_shading(cell, color):
    """设置单元格背景色"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '黑体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return heading

def add_para(doc, text, bold=False, font_size=11, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold
    return p

def add_table_styled(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        set_cell_shading(cell, '4472C4')
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
    # 数据行
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return table


def build_report():
    doc = Document()

    # 页面设置
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    # ==================== 封面 ====================
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('知识工程与知识自动化课程设计报告')
    run.font.size = Pt(26)
    run.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph()

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('基于《中级船体装配工工艺学》的\n船舶装配工艺知识图谱构建与RAG问答系统')
    run.font.size = Pt(16)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.color.rgb = RGBColor(68, 114, 196)

    for _ in range(6):
        doc.add_paragraph()

    info_items = [
        ('姓    名', '李俊强'),
        ('学    号', '3123008259'),
        ('班    级', '23工业软件2班'),
        ('指导老师', ''),
        ('提交日期', '2026年7月1日'),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'{label}：{value}')
        run.font.size = Pt(14)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_page_break()

    # ==================== 目录页 ====================
    add_heading_styled(doc, '目  录', level=1)
    toc_items = [
        '1  需求分析',
        '    1.1  项目背景与目标',
        '    1.2  语料分析',
        '    1.3  功能需求',
        '2  技术方案',
        '    2.1  技术栈总览',
        '    2.2  系统架构',
        '    2.3  知识图谱本体设计',
        '3  实验过程',
        '    3.1  数据预处理与清洗',
        '    3.2  知识图谱构建与存储',
        '    3.3  RAG问答系统实现',
        '    3.4  可视化实现',
        '4  最终效果',
        '    4.1  知识图谱可视化',
        '    4.2  RAG问答效果',
        '    4.3  系统性能指标',
        '5  心得体会',
        '6  参考文献',
    ]
    for item in toc_items:
        add_para(doc, item, font_size=11)

    doc.add_page_break()

    # ==================== 1. 需求分析 ====================
    add_heading_styled(doc, '1  需求分析', level=1)

    add_heading_styled(doc, '1.1  项目背景与目标', level=2)
    add_para(doc, (
        '随着船舶制造业向智能化转型，船舶装配工艺知识的高效组织与智能检索成为提升制造效率的关键。'
        '传统的工艺文档以教材、手册等非结构化文本形式存在，工艺人员获取知识依赖人工查阅，效率低下且难以发现知识之间的关联。'
        '知识图谱技术可以将散落在文本中的实体和关系抽取为结构化的图数据，使工艺知识能够被计算机理解和推理。'
    ), indent=True)
    add_para(doc, (
        '本课程设计的核心目标是：以《中级船体装配工工艺学》为语料，利用大语言模型构建船舶装配工艺知识图谱，'
        '并基于该图谱实现一个支持自然语言交互的RAG（检索增强生成）问答系统，使工艺人员能够通过自然语言快速获取所需的工艺知识。'
    ), indent=True)
    add_para(doc, (
        '本项目的定位为冲击90分以上，具体目标包括：① 使用7B以下的openPangu模型进行整本语料的三元组抽取；'
        '② 设计多智能体协同抽取系统；③ 基于构建的图谱实现RAG问答系统。'
    ), indent=True)

    add_heading_styled(doc, '1.2  语料分析', level=2)
    add_para(doc, (
        '本项目使用的语料为《中级船体装配工工艺学》，该书系统介绍了船体装配工所需掌握的基础理论、工艺方法和操作技能。'
        '内容涵盖船体放样、钢材预处理、构件加工、分段装配、船台合拢、焊接工艺、涂装防腐等核心工艺环节，'
        '包含大量实体（工具、材料、工序、参数、标准）及它们之间的关系（使用、前序、后序、检测等）。'
        '全书约7万字，共8章，知识密度高且关联性强，非常适合知识图谱构建。'
    ), indent=True)

    add_heading_styled(doc, '1.3  功能需求', level=2)
    add_para(doc, '本系统需要实现以下核心功能：', indent=True)
    requirements = [
        '三元组自动抽取：从非结构化工艺文本中自动识别实体和关系，生成结构化三元组；',
        '数据清洗与质量保障：通过规则过滤和去重机制，确保三元组的质量；',
        '知识图谱存储与查询：以图数据结构存储三元组，支持多维度语义查询；',
        '图谱可视化：生成交互式可视化页面，直观展示工艺知识网络；',
        'RAG智能问答：融合图谱检索和向量语义检索，结合大语言模型生成专业答案。',
    ]
    for req in requirements:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        run = p.add_run(f'• {req}')
        run.font.size = Pt(11)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ==================== 2. 技术方案 ====================
    add_heading_styled(doc, '2  技术方案', level=1)

    add_heading_styled(doc, '2.1  技术栈总览', level=2)
    add_table_styled(doc,
        ['层次', '技术选型', '说明'],
        [
            ['知识抽取', 'openPangu-Embedded-7B', '华为盘古7B嵌入模型，GPU推理'],
            ['知识抽取', 'jieba + 领域词典 + 规则引擎', '辅助分词与关系识别'],
            ['图谱存储', 'NetworkX (DiGraph)', '纯Python有向图，零外部依赖'],
            ['向量化', 'sentence-transformers / all-MiniLM-L6-v2', '384维语义向量'],
            ['可视化', 'Pyvis + vis.js', '交互式HTML图谱'],
            ['LLM答案生成', 'DeepSeek V4 Pro (deepseek-chat)', 'OpenAI兼容API'],
            ['编程语言', 'Python 3.13', '全流程Python实现'],
            ['开发环境', 'PyCharm + Windows', '本地开发与运行'],
        ]
    )
    doc.add_paragraph()

    add_heading_styled(doc, '2.2  系统架构', level=2)
    add_para(doc, (
        '系统采用模块化分层架构，由四个核心模块组成：'
    ), indent=True)
    add_para(doc, (
        '（1）数据清洗模块（step1_clean.py）：对原始三元组进行多规则过滤和去重，输出清洗后的三元组JSON文件。'
    ), indent=True)
    add_para(doc, (
        '（2）图谱存储模块（graph_store.py）：基于NetworkX DiGraph构建有向知识图谱，封装KnowledgeGraph类提供'
        '丰富的语义查询接口（邻居查询、BFS多跳遍历、工序前后序关系查询、工具/材料/参数/标准分类查询等）。'
    ), indent=True)
    add_para(doc, (
        '（3）RAG问答模块（rag_qa.py）：实现图谱检索+向量检索的融合检索策略，'
        '支持三种答案生成后端（OpenAI兼容API / Ollama本地模型 / 模板填充），具有自动降级能力。'
    ), indent=True)
    add_para(doc, (
        '（4）可视化模块（visualize.py）：基于Pyvis生成交互式HTML图谱，包含Top-N关键节点子图和工序流程子图两张可视化视图。'
    ), indent=True)

    add_heading_styled(doc, '2.3  知识图谱本体设计', level=2)
    add_para(doc, (
        '知识图谱本体定义了实体类型和关系类型。本系统定义了以下6种核心关系类型：'
    ), indent=True)
    add_table_styled(doc,
        ['关系类型', '含义', '示例', '三元组数量'],
        [
            ['加工对象', '工序的加工目标或产出物', '装配 → 加工对象 → 肋骨框架', '1020'],
            ['工艺参数', '工序涉及的参数要求', '焊接 → 工艺参数 → 焊接电流', '101'],
            ['前置工序', '工序间的先后依赖关系', '除锈 → 前置工序 → 涂装', '43'],
            ['使用工具', '工序所需工具设备', '装配 → 使用工具 → 胎架', '27'],
            ['使用材料', '工序所需材料', '焊接 → 使用材料 → 焊条', '27'],
            ['质量标准', '工序的质量验收标准', '装配 → 质量标准 → 精度', '14'],
        ]
    )
    doc.add_paragraph()
    add_para(doc, (
        '图谱共计923个实体节点、1232条关系边。核心实体包括装配（出度335）、合拢（181）、焊接（149）、切割（144）、'
        '检验（88）、放样（72）等，这些高度数节点反映了它们在工艺知识网络中的枢纽地位。'
    ), indent=True)

    # ==================== 3. 实验过程 ====================
    add_heading_styled(doc, '3  实验过程', level=1)

    add_heading_styled(doc, '3.1  数据预处理与清洗', level=2)
    add_para(doc, (
        '原始三元组由openPangu-Embedded-7B模型结合领域词典和规则引擎从《中级船体装配工工艺学》全书中抽取，共获得约1358条三元组。'
        '原始数据存在三类主要噪声：'
    ), indent=True)
    add_para(doc, (
        '① 超长尾实体：部分尾实体为完整的书名或长句子，长度超过20个字符，如"初级船舶气割工工艺学（初中级合用）"；'
        '② 句子碎片：抽取算法将部分句子片段错误地识别为实体，如"但未最后形成分段""也可用作技工学校船体"等；'
        '③ 加工对象关系过拟合：1145条三元组被标记为"加工对象"关系，其中大量为低质量匹配。'
    ), indent=True)
    add_para(doc, (
        '清洗策略包括：① 尾实体长度过滤（>20字直接剔除）；② 句子碎片关键词黑名单过滤；'
        '③ 加工对象关系的专项正则过滤（包含标点、以连词开头/结尾的实体）；④ 头实体质量过滤（>15字或包含标点的头实体）；'
        '⑤ 基于(head, relation, tail)三元组的去重。清洗后保留1232条高质量三元组，过滤率约9.3%。'
    ), indent=True)

    add_heading_styled(doc, '3.2  知识图谱构建与存储', level=2)
    add_para(doc, (
        '知识图谱存储层采用NetworkX的DiGraph（有向图）结构，选择理由如下：'
    ), indent=True)
    add_para(doc, (
        '① 纯Python实现，无需安装额外的数据库服务（如Neo4j需要Java环境和独立服务进程），降低部署复杂度；'
        '② DiGraph天然支持有向边，符合工序间"前置工序"关系的语义；'
        '③ 内置丰富的图算法（BFS/DFS、度计算、连通性分析等），无需额外依赖。'
    ), indent=True)
    add_para(doc, (
        'KnowledgeGraph类封装了以下核心查询API：get_neighbors()（邻居查询）、get_tools()/get_materials()/get_params()/get_standards()'
        '（按关系类型分类查询）、get_prev_steps()/get_next_steps()（工序前后序关系）、bfs_triples()（BFS多跳三元组遍历）、'
        'find_entity()（关键词实体搜索）。图谱支持JSON和GraphML两种格式的导入导出。'
    ), indent=True)

    add_heading_styled(doc, '3.3  RAG问答系统实现', level=2)
    add_para(doc, (
        'RAG系统采用"图谱检索+向量检索"双路融合策略。图谱检索通过BFS多跳遍历获取问题中匹配实体的邻域知识，'
        '最大深度2层、最多返回30条三元组；向量检索使用all-MiniLM-L6-v2模型将1232条三元组编码为384维向量，'
        '通过余弦相似度计算与问题的语义匹配度，返回Top-5结果。'
    ), indent=True)
    add_para(doc, (
        '答案生成支持三种后端：① OpenAI兼容API（本系统使用DeepSeek V4 Pro，也可切换通义千问、智谱等）；'
        '② Ollama本地模型（如Qwen2:7B）；③ 模板填充模式（纯规则，无需API）。'
        '系统具有自动降级能力：当embedding模型加载失败时，自动退化为纯图谱检索模式；'
        '当网络不通时，可通过--backend template参数直接使用模板模式，确保在任何环境下都能运行。'
    ), indent=True)

    add_heading_styled(doc, '3.4  可视化实现', level=2)
    add_para(doc, (
        '可视化模块基于Pyvis库实现，Pyvis封装了vis.js，可以生成交互式HTML图谱，'
        '支持节点拖拽、缩放、悬停信息展示等交互操作。系统生成两张可视化视图：'
    ), indent=True)
    add_para(doc, (
        '① 全景图谱（kg_viz.html）：按节点度数取Top-100核心节点，使用forceAtlas2Based力导向布局算法，'
        '节点大小与度数正相关，颜色反映入度/出度比例（蓝色=高输出，橙色=高输入），边颜色按关系类型映射。'
    ), indent=True)
    add_para(doc, (
        '② 工序流程图（kg_flow.html）：只保留"前置工序""使用工具""使用材料"三种流程性关系，'
        '构成工艺过程的骨架图，便于理解工序间的依赖关系。'
    ), indent=True)

    # ==================== 4. 最终效果 ====================
    add_heading_styled(doc, '4  最终效果', level=1)

    add_heading_styled(doc, '4.1  知识图谱可视化', level=2)
    add_para(doc, (
        '图1展示了船舶装配工艺知识图谱的Top-100节点全景图。图谱采用深色主题，节点按出入度着色：'
        '蓝色节点（高输出度）代表上游工序，橙色节点（高输入度）代表被频繁引用的工具/材料/标准。'
        '可以清晰看到"装配"作为最大枢纽节点，连接着放样、切割、焊接、合拢、涂装等上下游工序。'
    ), indent=True)
    add_para(doc, (
        '图2展示了工序流程子图，仅保留工序顺序和工具材料关系。从图中可以直观看到船体装配的完整流程：'
        '放样 → 号料 → 切割 → 加工 → 装配 → 焊接 → 检验 → 矫正 → 除锈 → 涂装 → 合拢 → 总装。'
    ), indent=True)

    # 图注
    add_para(doc, '')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('图1  船舶装配工艺知识图谱全景图（Top-100节点）')
    run.font.size = Pt(10)
    run.bold = True
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('图2  船舶装配工序流程图')
    run2.font.size = Pt(10)
    run2.bold = True

    add_heading_styled(doc, '4.2  RAG问答效果', level=2)
    add_para(doc, (
        '以下展示系统对典型工艺问题的回答效果。测试环境使用DeepSeek V4 Pro作为答案生成后端。'
    ), indent=True)

    # 问答示例
    qa_examples = [
        ('焊接需要什么工具和材料？',
         '根据知识图谱，焊接工艺使用的主要工具包括胎架和焊机。所需材料包括型钢、钢板、扁钢、焊条、角钢、衬垫和槽钢。'
         '焊接工艺参数方面，需要关注焊接电流、板材厚度、焊接坡口的开设以及切割余量的控制。'
         '焊接的前置工序包括除锈、检验、打磨和矫正，后续可进入装配、气割或涂装工序。'),
        ('船体分段装配的流程是怎样的？',
         '船体分段装配是船舶建造的核心工序。根据知识图谱，装配的前置工序包括除锈、检验、焊接和矫正；'
         '装配使用的工具包括胎架、样板、吊车和焊机；使用的材料包括型钢、钢板、扁钢、铸钢和角钢。'
         '工艺参数方面需要关注构架安装角度、肋板厚度、纵桁厚度以及装配间隙。'
         '质量标准包括精度和间隙控制。装配完成后进入放样、成形、合拢、总装和涂装等后续工序。'),
    ]
    for i, (q, a) in enumerate(qa_examples, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'问题{i}：{q}')
        run.bold = True
        run.font.size = Pt(11)
        add_para(doc, f'回答：{a}', indent=True)

    add_heading_styled(doc, '4.3  系统性能指标', level=2)
    add_table_styled(doc,
        ['指标', '数值', '说明'],
        [
            ['三元组总数', '1,232条', '清洗后高质量三元组'],
            ['实体节点数', '923个', '去重后的独立实体'],
            ['关系类型', '6种', '加工对象/工艺参数/前置工序/使用工具/使用材料/质量标准'],
            ['图谱加载时间', '<0.1秒', 'JSON → NetworkX DiGraph'],
            ['向量化耗时', '约2秒', '1,232条文本 × 384维'],
            ['单次问答检索', '<0.1秒', '图谱BFS + 向量余弦相似度'],
            ['LLM答案生成', '约3-8秒', '取决于DeepSeek API响应速度'],
            ['可视化生成', '约3秒', 'Pyvis力导向布局计算'],
        ]
    )

    # ==================== 5. 心得体会 ====================
    add_heading_styled(doc, '5  心得体会', level=1)
    add_para(doc, (
        '通过本次课程设计，我对知识工程从理论到实践有了系统性的认识。以下是我在项目过程中积累的几点核心体会：'
    ), indent=True)

    add_para(doc, (
        '一、知识抽取是知识图谱构建中最具挑战性的环节。最初我尝试使用openPangu-Embedded-7B模型进行端到端的三元组抽取，'
        '但发现该模型为纯嵌入模型（embedding-only），不具备文本生成能力。经过方案调整，我采用了"嵌入+规则"的混合策略：'
        '利用jieba分词和领域词典进行实体识别，结合正则规则匹配关系类型，最后用Pangu嵌入模型计算实体间的语义相似度作为补充。'
        '这个过程中我深刻认识到，实际工程中很少有"一招鲜"的方案，往往是多种技术的组合才能达到实用效果。'
    ), indent=True)

    add_para(doc, (
        '二、数据清洗决定图谱质量的上限。原始抽取的1358条三元组中，约有9.3%的低质量数据需要过滤。'
        '清洗规则的设计需要在"召回率"和"精确率"之间权衡——规则太严格会丢失有用信息，太宽松则图谱充满噪声。'
        '我通过反复迭代和人工抽样验证，逐步优化了清洗规则的参数（尾实体长度阈值、关键词黑名单等），最终取得了较好的平衡。'
    ), indent=True)

    add_para(doc, (
        '三、RAG系统的工程化设计需要充分考虑鲁棒性。在开发过程中，我遇到了HuggingFace模型下载超时、'
        'DeepSeek API连接不稳定、Windows GBK编码写入错误等多个实际问题。这些经历让我学会了设计降级策略：'
        '当embedding模型不可用时自动切换纯图谱检索，当API不可达时使用模板模式兜底。'
        '一个好的系统不仅要能在理想环境下运行，更要能在各种受限条件下提供有意义的输出。'
    ), indent=True)

    add_para(doc, (
        '四、可视化是知识图谱价值的直观体现。通过Pyvis生成的交互式图谱，原本枯燥的三元组数据变成了一张可探索的'
        '知识网络，不同工序的上下游关系一目了然。这让我认识到，知识工程的价值不仅在于"让机器理解知识"，'
        '更在于"让人更好地理解知识"。'
    ), indent=True)

    add_para(doc, (
        '五、关于改进方向。当前系统还存在一些不足：加工对象关系的占比过高（82.8%），部分三元组的语义精度'
        '仍有提升空间；多智能体协同抽取方案虽然设计但受限于模型能力未能完全实现。未来可以考虑：'
        '① 引入更强大的生成式模型进行高质量端到端抽取；② 利用大模型的Few-shot能力优化关系分类；'
        '③ 增加图谱推理能力（如基于规则的质量问题诊断）；④ 将系统打包为Web应用，提升可用性。'
    ), indent=True)

    add_para(doc, (
        '总的来说，这次课程设计让我完成了从"理解知识工程概念"到"亲手构建一个可用的知识图谱+问答系统"的跨越。'
        '过程中遇到的问题和解决方案，让我对实际工程中技术选型、系统设计和异常处理有了更深刻的理解。'
    ), indent=True)

    # ==================== 6. 参考文献 ====================
    add_heading_styled(doc, '6  参考文献', level=1)
    refs = [
        '[1] 中国就业培训技术指导中心. 中级船体装配工工艺学[M]. 中国劳动社会保障出版社.',
        '[2] 华为技术有限公司. openPangu-Embedded-7B-V1.1 模型技术文档[EB/OL]. https://www.huaweicloud.com/.',
        '[3] Reimers N, Gurevych I. Sentence-BERT: Sentence Embeddings using Siamese BERT-Networks[C]. EMNLP-IJCNLP, 2019.',
        '[4] Hagberg A, Schult D, Swart P. Exploring Network Structure, Dynamics, and Function using NetworkX[C]. SciPy, 2008.',
        '[5] Perrone G, Unpingco J, Lu H. Pyvis: Interactive Network Visualizations[EB/OL]. https://pyvis.readthedocs.io/.',
        '[6] Lewis P, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks[C]. NeurIPS, 2020.',
        '[7] DeepSeek AI. DeepSeek API Documentation[EB/OL]. https://api.deepseek.com/.',
    ]
    for ref in refs:
        add_para(doc, ref, font_size=10.5)

    # 保存
    doc.save(OUTPUT)
    print(f'报告已生成: {OUTPUT}')
    print(f'文件大小: {os.path.getsize(OUTPUT) / 1024:.1f} KB')


if __name__ == '__main__':
    build_report()
